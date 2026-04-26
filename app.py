import os
import io
from pathlib import Path
from openai import OpenAI
from flask import Flask, render_template, request, jsonify, send_file
from dotenv import load_dotenv

load_dotenv(Path(__file__).parent / '.env')

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", os.urandom(24))
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 最大上传 5MB

client = OpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY"),
    base_url="https://api.deepseek.com",
)

# ── 工具函数 ──────────────────────────────────────────────────

def extract_text_from_file(file) -> str:
    """从上传的文件中提取文本（支持 .txt / .pdf / .docx）"""
    filename = file.filename.lower()
    content = file.read()

    if filename.endswith('.txt'):
        return content.decode('utf-8', errors='ignore')

    if filename.endswith('.pdf'):
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                return '\n'.join(p.extract_text() or '' for p in pdf.pages)
        except Exception as e:
            return f"[PDF 解析失败: {e}]"

    if filename.endswith('.docx'):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            return '\n'.join(p.text for p in doc.paragraphs)
        except Exception as e:
            return f"[DOCX 解析失败: {e}]"

    return "[不支持的文件格式，请上传 .txt / .pdf / .docx]"


def call_deepseek(system_prompt: str, user_prompt: str, max_tokens=2000) -> str:
    """统一调用 DeepSeek"""
    response = client.chat.completions.create(
        model="deepseek-chat",
        max_tokens=max_tokens,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
    )
    return response.choices[0].message.content.strip()


def generate_resume_html(data: dict) -> str:
    """根据表单数据用 AI 生成简历 HTML 片段（只有内容，不含页面框架）"""
    system = """你是一位专业的简历撰写专家。
根据用户提供的信息，生成一份结构清晰、语言专业的简历内容。
输出格式：纯 HTML 片段，不要包含 <html>/<head>/<body> 标签。
使用以下结构（每个 section 按需生成，没有内容则跳过）：
- <section class="resume-section" id="profile"> 个人信息
- <section class="resume-section" id="summary"> 个人简介
- <section class="resume-section" id="experience"> 工作经历
- <section class="resume-section" id="education"> 教育背景
- <section class="resume-section" id="skills"> 技能
- <section class="resume-section" id="projects"> 项目经历（可选）
每段工作/教育使用 <div class="resume-item"> 包裹。
语言风格：简洁、有力、专业。"""

    user = f"""请根据以下信息生成一份完整的简历：

姓名：{data.get('name', '')}
求职岗位：{data.get('target_job', '')}
联系方式：{data.get('contact', '')}
个人简介：{data.get('summary', '')}
工作经历：{data.get('experience', '')}
教育背景：{data.get('education', '')}
技能：{data.get('skills', '')}
项目经历：{data.get('projects', '')}"""

    return call_deepseek(system, user)


def polish_resume(old_text: str, target_job: str) -> str:
    """对旧简历进行 AI 润色，返回 HTML 片段"""
    system = """你是一位专业的简历优化专家。
对用户提供的旧简历进行全面润色：
1. 优化措辞，使语言更专业有力
2. 突出与目标岗位相关的经历和技能
3. 调整结构，使简历更易阅读
输出格式：纯 HTML 片段，结构同上（使用 section.resume-section + div.resume-item）。"""

    user = f"目标岗位：{target_job}\n\n旧简历内容：\n{old_text}"
    return call_deepseek(system, user)


def html_to_docx(html_content: str, candidate_name: str) -> bytes:
    """将 AI 生成的简历 HTML 解析后写入 Word 文档，返回字节流"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()

    # ── 页边距 ──
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── 默认正文样式 ──
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def add_separator(doc):
        """添加细横线分隔符"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2563EB')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_section_heading(doc, text):
        """蓝色加粗节标题 + 下划线"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        add_separator(doc)

    def add_item_header(doc, title, meta=''):
        """条目标题行：加粗标题 + 右侧日期/机构"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if meta:
            run2 = p.add_run(f'  {meta}')
            run2.font.size = Pt(9)
            run2.font.color.rgb = RGBColor(0x55, 0x60, 0x70)
            run2.font.name = '微软雅黑'
            run2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def add_body_text(doc, text):
        if not text.strip():
            return
        p = doc.add_paragraph(text.strip())
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def add_bullet(doc, text):
        if not text.strip():
            return
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(text.strip())
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ── 遍历各 section ──
    for section_el in soup.find_all('section', class_='resume-section'):
        sec_id = section_el.get('id', '')

        # 个人信息区：姓名居中大字
        if sec_id == 'profile':
            h1 = section_el.find('h1')
            if h1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(h1.get_text(strip=True))
                run.bold = True
                run.font.size = Pt(20)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            for tag in section_el.find_all(['p', 'div']):
                txt = tag.get_text(' | ', strip=True)
                if txt:
                    p2 = doc.add_paragraph(txt)
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p2.paragraph_format.space_after = Pt(4)
                    for run in p2.runs:
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0x55, 0x60, 0x70)
            continue

        # 节标题
        h2 = section_el.find('h2')
        if h2:
            add_section_heading(doc, h2.get_text(strip=True))

        # 条目
        items = section_el.find_all('div', class_='resume-item')
        if items:
            for item in items:
                h3 = item.find('h3')
                meta = item.find(class_='meta')
                title_text = h3.get_text(strip=True) if h3 else ''
                meta_text  = meta.get_text(strip=True) if meta else ''
                if title_text:
                    add_item_header(doc, title_text, meta_text)
                for li in item.find_all('li'):
                    add_bullet(doc, li.get_text(strip=True))
                for p_tag in item.find_all('p'):
                    add_body_text(doc, p_tag.get_text(strip=True))
        else:
            # 没有 resume-item 的节（如技能）直接输出文本/列表
            for li in section_el.find_all('li'):
                add_bullet(doc, li.get_text(strip=True))
            for p_tag in section_el.find_all('p'):
                add_body_text(doc, p_tag.get_text(strip=True))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── 路由 ──────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/generate', methods=['POST'])
def generate():
    """从表单数据生成简历"""
    try:
        data = {
            'name':       request.form.get('name', '').strip(),
            'target_job': request.form.get('target_job', '').strip(),
            'contact':    request.form.get('contact', '').strip(),
            'summary':    request.form.get('summary', '').strip(),
            'experience': request.form.get('experience', '').strip(),
            'education':  request.form.get('education', '').strip(),
            'skills':     request.form.get('skills', '').strip(),
            'projects':   request.form.get('projects', '').strip(),
        }
        if not data['name']:
            return jsonify({'error': '姓名不能为空'}), 400

        resume_html = generate_resume_html(data)
        return jsonify({'html': resume_html, 'name': data['name']})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/polish', methods=['POST'])
def polish():
    """上传旧简历并润色"""
    try:
        target_job = request.form.get('target_job', '').strip()
        file = request.files.get('resume_file')

        if not file or file.filename == '':
            return jsonify({'error': '请上传简历文件'}), 400

        old_text = extract_text_from_file(file)
        if old_text.startswith('['):
            return jsonify({'error': old_text}), 400

        resume_html = polish_resume(old_text, target_job)
        return jsonify({'html': resume_html})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download-word', methods=['POST'])
def download_word():
    """将简历 HTML 转换为 Word 文档并下载"""
    try:
        resume_html = request.json.get('html', '')
        name = request.json.get('name', '简历')
        if not resume_html:
            return jsonify({'error': '内容为空'}), 400

        docx_bytes = html_to_docx(resume_html, name)
        return send_file(
            io.BytesIO(docx_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{name}_简历.docx'
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=os.getenv("FLASK_DEBUG", "false").lower() == "true")